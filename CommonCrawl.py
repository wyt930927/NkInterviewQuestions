import urllib
import docx
import re
import datetime
from bs4 import BeautifulSoup

urls = ['/ta/review-frontend', '/ta/review-java', '/ta/review-network', '/ta/front-end-interview', '/ta/nine-chapter']
titles = ['前端面试常考HTML+CSS', 'JAVA面试常考知识点', '计算机网络面试常考点', 'JS面试经典题目合集', 'BAT经典面试题汇总']
authority = 'https://www.nowcoder.com'


def __parserHtml__(soup, url1, index):
    print(titles[index])
    file = docx.Document()
    i = 0
    pul = soup.select('.pagination ul')
    if len(pul) > 0:
        page_total = int(pul[0]['data-total'])
        for page_now in range(1, page_total):
            if page_now != 1:
                soup = __craw_html__(url1 + '?query=&asc=true&order=&page=' + str(page_now))
            t_urls = soup.findAll('table')[0]
            for tr in t_urls.tbody.findAll('tr'):
                o = tr.select('td:nth-of-type(2)')
                if len(o) > 0:
                    tag = o[0].find('a')
                    i += 1
                    href = tag['href']
                    soup1 = __craw_html__(authority + href)
                    file.add_paragraph(str(i) + '.' + soup1.select('.final-question')[0].text.replace('\n', ''))
                    print(str(i) + '.' + tag.text)
                    answers = soup1.select('.design-answer-box')[0]
                    if len(answers.select('p')) > 0:
                        for txt in answers.select('p'):
                            line = re.sub('<[^<]+?>', '', txt.text).strip()
                            file.add_paragraph(line)
                    else:
                        file.add_paragraph(re.sub('<[^<]+?>', '', answers.text.strip()))
                        # print(re.sub('<[^<]+?>', '', answers.text).replace('\n', '').strip())
    file.save("D:\\temp\\" + titles[index] + ".docx")


def __craw_html__(url1):
    res = urllib.request.urlopen(url1)
    html = res.read().decode('utf-8')
    return BeautifulSoup(html, 'html.parser')


if __name__ == '__main__':
    start_time = datetime.datetime.now()
    for ind in range(len(urls)):
        url0 = authority + urls[ind]
        soup0 = __craw_html__(url0)
        __parserHtml__(soup0, url0, ind)
    end_time = datetime.datetime.now()
    print('执行结束，共计' + str((end_time - start_time).seconds) + 's')
